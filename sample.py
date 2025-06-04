from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
from io import BytesIO

def add_html_with_images(doc, html_content, local_image_path):
    soup = BeautifulSoup(html_content, 'html.parser')

    for elem in soup.contents:
        if elem.name == 'img':
            try:
                with open(local_image_path, 'rb') as f:
                    img_stream = BytesIO(f.read())
                    doc.add_picture(img_stream, width=Inches(4))
            except Exception as e:
                doc.add_paragraph(f"[Image failed to load from {local_image_path}: {e}]")
        elif elem.name:
            text = elem.get_text(strip=True)
            if text:
                doc.add_paragraph(text)
        else:
            text = str(elem).strip()
            if text:
                doc.add_paragraph(text)


doc = Document()
local_image_path = r"D:\servicenow_api_ops\KB_docx_files_20250604_075811\KB0019598\Pasted image.png"
html_content = '''
    <p>Some text here</p>
    <img src="anything.jpg" />
    <p>More text here</p>
'''
add_html_with_images(doc, html_content, local_image_path)

doc.save("output.docx")
