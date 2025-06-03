from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
import requests
from io import BytesIO
from dotenv import load_dotenv
import os 
from urllib.parse import urlencode

load_dotenv()

html_content = '''<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Wellbeing leave is one of the ways Lendlease supports employee&rsquo;s physical and mental wellbeing. This leave is provided to eligible employees in addition to standard leave entitlements.</span></p>\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Wellbeing Leave can only be taken as a full or half day, per leave application. A maximum of one day's Wellbeing Leave can be taken in each specified four month period, as follows:</span></p>\r\n<ul style=\"list-style-position: inside;\">\r\n<li><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">1 July to 31 October</span></li>\r\n<li><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">1 November to 28 February (or 29 February in a leap year); and</span></li>\r\n<li><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">1 March to 30 June.</span></li>\r\n</ul>\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Wellbeing Leave is not accrued and does not carry over from one period to the next. Any Wellbeing Leave days not taken during a specified four month period will lapse.</span></p>\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Employees who commence employment anytime during a specified four month period will be eligible for Wellbeing Leave at the start of the next period.</span></p>\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Similarly, employees with excessive Annual Leave balances (30 days or more) at the time of application are not eligible, until they reduce their balance to below this threshold.</span></p>\r\n<p><br /></p>\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\"><strong>What are the eligibility requirements?</strong></span></p>\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Permanent and fixed-term salaried employees of Lendlease are eligible for wellbeing leave, subject to policy conditions.</span><br /><br /><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Where employees are employed under the terms and conditions of an operative enterprise agreement with a Lendlease group company, please refer to the relevant enterprise agreement as to whether or not this policy applies to you.</span><br /><br /><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Additionally, employees with excessive annual leave balances of 30 days or more are not eligible.</span><br /><br /><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Employees who start part way through a specified four-month period (per the policy) will be eligible for Wellbeing Leave at the start of the next period.</span></p>\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\"><strong>How do I apply for wellbeing leave?</strong></span></p>\r\n<table style=\"width: auto;\">\r\n<tbody>\r\n<tr style=\"height: 19px;\">\r\n<td style=\"width: auto; background-color: #ffffff; border-right: 1.5pt solid #5c5858; padding: 6.75pt 7.5pt; border-top-style: solid; border-bottom-style: solid; border-left-style: solid; border-top-color: #5c5858; border-bottom-color: #5c5858; border-left-color: #5c5858;\">\r\n<p><strong><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">PLAN</span></strong></p>\r\n</td>\r\n<td style=\"width: auto; background-color: #ffffff; border-right: 1.5pt solid #5c5858; padding: 6.75pt 7.5pt; border-top-style: solid; border-bottom-style: solid; border-left-style: solid; border-top-color: #5c5858; border-bottom-color: #5c5858; border-left-color: #5c5858;\">\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Read through the Wellbeing Leave Policy.&nbsp;</span></p>\r\n</td>\r\n</tr>\r\n<tr style=\"height: 19px;\">\r\n<td style=\"width: auto; background-color: #ffffff; border-right: 1.5pt solid #5c5858; padding: 6.75pt 7.5pt; border-top-style: solid; border-bottom-style: solid; border-left-style: solid; border-top-color: #5c5858; border-bottom-color: #5c5858; border-left-color: #5c5858;\">\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\"><strong>PREPARE</strong>&nbsp;</span></p>\r\n</td>\r\n<td style=\"width: auto; background-color: #ffffff; border-right: 1.5pt solid #5c5858; padding: 6.75pt 7.5pt; border-top-style: solid; border-bottom-style: solid; border-left-style: solid; border-top-color: #5c5858; border-bottom-color: #5c5858; border-left-color: #5c5858;\">\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Discuss your leave plans with your manager well in advance, in order to agree to a time which minimises disruption to the business.&nbsp;</span></p>\r\n</td>\r\n</tr>\r\n<tr style=\"height: 19px;\">\r\n<td style=\"width: auto; background-color: #ffffff; border-right: 1.5pt solid #5c5858; padding: 6.75pt 7.5pt; border-top-style: solid; border-bottom-style: solid; border-left-style: solid; border-top-color: #5c5858; border-bottom-color: #5c5858; border-left-color: #5c5858;\">\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\"><strong>ACT</strong>&nbsp;</span></p>\r\n</td>\r\n<td style=\"width: auto; background-color: #ffffff; border-right: 1.5pt solid #5c5858; padding: 6.75pt 7.5pt; border-top-style: solid; border-bottom-style: solid; border-left-style: solid; border-top-color: #5c5858; border-bottom-color: #5c5858; border-left-color: #5c5858;\">\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Once you have agreed your leave plans with your manager, submit a leave request in Workday.&nbsp;</span></p>\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Your Wellbeing leave request will be automatically approved in Workday.</span></p>\r\n</td>\r\n</tr>\r\n<tr style=\"height: 19px;\">\r\n<td style=\"width: auto; background-color: #ffffff; border-right: 1.5pt solid #5c5858; padding: 6.75pt 7.5pt; border-top-style: solid; border-bottom-style: solid; border-left-style: solid; border-top-color: #5c5858; border-bottom-color: #5c5858; border-left-color: #5c5858;\">\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\"><strong>FINALISE</strong>&nbsp;</span></p>\r\n</td>\r\n<td style=\"width: auto; background-color: #ffffff; border-right: 1.5pt solid #5c5858; padding: 6.75pt 7.5pt; border-top-style: solid; border-bottom-style: solid; border-left-style: solid; border-top-color: #5c5858; border-bottom-color: #5c5858; border-left-color: #5c5858;\">\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">If necessary, create or revise your department calendars to reflect your time away from the office and/or out of office automatic replies within MSOffice outlook.</span></p>\r\n</td>\r\n</tr>\r\n</tbody>\r\n</table>\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">How do I cancel or amend my wellbeing leave request?</span></p>\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">You can cancel and amend leave requests through Workday. Where possible, you should discuss any changes to leave plans with your manager prior to amending in Workday.</span></p>\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\"><strong>Resources</strong></span></p>\r\n<table style=\"width: auto;\">\r\n<tbody>\r\n<tr style=\"height: 19px;\">\r\n<td style=\"width: 60.0pt; background-color: #ffffff; vertical-align: center; border-right: 1.5pt solid #FFFFFF; padding: 6.75pt 7.5pt 6.75pt 7.5pt;\">\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">&nbsp;<img src=\"/sys_attachment.do?sys_id=481d5d331bbba010ef4543b4bd4bcb0a\" alt=\"Stationery_Paper_BLK_PNG png\" width=\"48px\" height=\"59px\" /></span></p>\r\n</td>\r\n<td style=\"width: 525.0pt; background-color: #ffffff; vertical-align: center; border-right: 1.5pt solid #FFFFFF; padding: 6.75pt 7.5pt 6.75pt 7.5pt;\">\r\n<p><a href=\"https://lendlease.service-now.com/lendlease?id=kb_article_view&amp;sysparm_article=KB0018962\" target=\"_blank\" rel=\"noopener\" title=\"Wellbeing Leave Policy\"><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">Wellbeing Leave Policy</span></a></p>\r\n</td>\r\n</tr>\r\n<tr style=\"height: 19px;\">\r\n<td style=\"width: 60.0pt; background-color: #ffffff; vertical-align: center; border-right: 1.5pt solid #FFFFFF; padding: 6.75pt 7.5pt 6.75pt 7.5pt;\">\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\">&nbsp;<img src=\"/sys_attachment.do?sys_id=081ddd331bbba010ef4543b4bd4bcb62\" alt=\"Resources_BLK_PNG png\" width=\"48px\" height=\"45px\" /></span></p>\r\n</td>\r\n<td style=\"width: auto; background-color: #ffffff; vertical-align: center; border-right: 1.5pt solid #FFFFFF; padding: 6.75pt 7.5pt 6.75pt 7.5pt;\">\r\n<p><span style=\"font-family: arial, helvetica, sans-serif; font-size: 12pt;\"><a href=\"https://wd3.myworkday.com/lendlease/login.htmld\">Workday</a></span></p>\r\n</td>\r\n</tr>\r\n</tbody>\r\n</table>'''
def get_bearer_token():
    url = "https://lendlease.service-now.com/oauth_token.do"
    payload_dict = {
        'grant_type': 'password',
        'username': os.getenv('SNOW_USERNAME'),
        'password': os.getenv('SNOW_PASSWORD'),
        'client_id': os.getenv('SNOW_CLIENT_ID'),
        'client_secret': os.getenv('SNOW_CLIENT_SECRET')
    }
    payload = urlencode(payload_dict)

    headers = {
    'Content-Type': 'application/x-www-form-urlencoded',
    }
    response = requests.post(url, data=payload, headers=headers)
    print(f"Response Status Code: {response.status_code}")
    data = response.json()
    return data['access_token']

base_url = "https://lendlease.service-now.com"
token = get_bearer_token()
headers = {
  'Authorization': f'Bearer {token}',
  'Cookie': 'BIGipServerpool_lendlease=c5889ad29f701618e3baa37002034b82; JSESSIONID=3901AC59B602B51CE1CF74C8956FD362; glide_node_id_for_js=fc4812175032dd94c0ff92cf846b17cf27f0dce0a6beb49e12e5c7bb0f48d836; glide_session_store=6360D6592B3D6E50E412F41CD891BF5D; glide_user_activity=U0N2M18xOnRMdkppdFlTN2o2cFlnUVdaQ092UjZ6S0pFdXV0dmZBb3BMcGxVa0hrZ1E9OlVBQWc4QWozUERYQi9mVCs2WDRJa0hTRTgwQjkxMGZkMzUrNGxlUXRNUW89; glide_user_route=glide.5a07cc0a1b859ed021434a69d48daaeb'
}


doc = Document()
soup = BeautifulSoup(html_content, "html.parser")

for element in soup.children:
    if element.name == "p":
        # Add paragraph text (with inner text only)
        text = element.get_text(strip=True)
        if text:
            doc.add_paragraph(text)
    elif element.name == "ul":
        # Add bullet points
        for li in element.find_all("li"):
            doc.add_paragraph(li.get_text(strip=True), style="List Bullet")
    elif element.name == "table":
        # Handle tables if needed - simplified here
        table = doc.add_table(rows=0, cols=2)
        for tr in element.find_all("tr"):
            cells = tr.find_all(["td", "th"])
            row_cells = table.add_row().cells
            for i in range(min(2, len(cells))):
                row_cells[i].text = cells[i].get_text(strip=True)
    elif element.name == "img":
        # Handle any standalone images
        img_src = element.get("src")
        if img_src:
            if img_src.startswith("/"):
                img_src = base_url + img_src
            try:
                # response = requests.get(img_src)
                response = requests.get(img_src, headers=headers)
                response.raise_for_status()
                image_stream = BytesIO(response.content)
                doc.add_picture(image_stream, width=Inches(2))
            except Exception as e:
                print(f"Failed to download image {img_src}: {e}")

# Additionally, handle images embedded inside paragraphs or other tags
for img_tag in soup.find_all("img"):
    img_src = img_tag.get("src")
    if img_src:
        if img_src.startswith("/"):
            img_src = base_url + img_src
        try:
            # response = requests.get(img_src)
            response = requests.get(img_src, headers=headers)
            response.raise_for_status()
            image_stream = BytesIO(response.content)
            doc.add_picture(image_stream, width=Inches(2))
        except Exception as e:
            print(f"Failed to download image {img_src}: {e}")

# Save doc
doc.save("wellbeing_leave.docx")