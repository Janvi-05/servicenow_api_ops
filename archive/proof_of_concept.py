from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup, NavigableString, Tag

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

def add_hyperlink(paragraph, url, text, styles=None):
    if styles is None:
        styles = {}

    # Create the hyperlink relationship
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a run for the hyperlink
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Apply hyperlink style
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    # Apply additional styles
    if styles.get("bold"):
        b = OxmlElement("w:b")
        rPr.append(b)
    if styles.get("italic"):
        i = OxmlElement("w:i")
        rPr.append(i)
    if styles.get("strike"):
        strike = OxmlElement("w:strike")
        rPr.append(strike)

    new_run.append(rPr)

    # Add text
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def process_node(paragraph, node, inherited_styles=None):
    if inherited_styles is None:
        inherited_styles = {}

    if isinstance(node, NavigableString):
        run = paragraph.add_run(str(node))
        if inherited_styles.get('bold'):
            run.bold = True
        if inherited_styles.get('italic'):
            run.italic = True
        if inherited_styles.get('strike'):
            run.font.strike = True
    elif isinstance(node, Tag):
        new_styles = inherited_styles.copy()
        if node.name in ['b', 'strong']:
            new_styles['bold'] = True
        if node.name in ['i', 'em']:
            new_styles['italic'] = True
        if node.name in ['s', 'strike', 'del']:
            new_styles['strike'] = True

        if node.name == 'a' and node.get('href'):
            add_hyperlink(paragraph, node['href'], node.get_text(), styles=new_styles)
        else:
            for child in node.children:
                process_node(paragraph, child, new_styles)

def html_to_docx(html: str, output_path: str):
    doc = Document()
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.find_all(['p', 'div', 'span']):
        paragraph = doc.add_paragraph()
        process_node(paragraph, element)

    doc.save(output_path)

html_content = """
<p>This is a <b>bold</b>, <i>italic</i>, and <s>strike-through</s> text with a 
<b><i><a href="https://example.com">styled hyperlink</a></i></b>.</p>
"""
html_to_docx(html_content, 'output.docx')

